"""动态 Word 渲染引擎 —— 基于模板样式 + HierarchyProfile，将 JSON blocks 渲染为 .docx。

工作流程：加载模板 → 清空原有内容 → 遍历 JSON blocks → 按 type/level 动态排版。

排版规则（对齐示例文档，由 HierarchyProfile 驱动）：
  - title / heading: Heading {level} 样式 + profile.style_map[level] run 级字体覆写
  - annotation cn_subtitle: Heading {level} 样式 + profile.style_map[level] run 级字体覆写
  - annotation en_text: Normal 样式 + profile.en_text_style run 级字体覆写
"""

from __future__ import annotations

import logging
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

if TYPE_CHECKING:
    from hierarchy_analyzer import HierarchyProfile

logger = logging.getLogger(__name__)

_ILLEGAL_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _sanitize(text: str) -> str:
    """移除 XML 不兼容的控制字符，避免 python-docx 写入时崩溃。"""
    return _ILLEGAL_XML_RE.sub("", text)


def _set_outline_level(paragraph, level: int):
    """强制写入 w:outlineLvl 到段落 pPr，确保 WPS 可识别大纲层级。

    level 使用 1-based（Heading 1 = 1），内部转为 0-based 写入 XML。
    """
    pPr = paragraph._element.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = paragraph._element.makeelement(qn("w:outlineLvl"), {})
        pPr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def _set_outline_level_body(paragraph):
    """强制写入 outlineLvl=9（正文级别）到段落 pPr，明确标记为非标题。"""
    pPr = paragraph._element.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = paragraph._element.makeelement(qn("w:outlineLvl"), {})
        pPr.append(outline)
    outline.set(qn("w:val"), "9")


def _set_run_fonts(
    run,
    latin_font: str = "Times New Roman",
    east_asia_font: str | None = None,
    size_pt: float = 11,
    bold: bool = False,
):
    """统一设置 run 的中英文字体、字号和加粗。"""
    run.font.name = latin_font
    run.font.size = Pt(size_pt)
    run.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    if east_asia_font:
        rFonts.set(qn("w:eastAsia"), east_asia_font)


def _set_paragraph_spacing(paragraph, line_spacing: float | None):
    """设置段落行距（仅当 line_spacing 非 None 时覆写）。"""
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing


def _apply_profile_fonts(
    run, paragraph, level: int,
    profile: HierarchyProfile | None,
):
    """根据 profile.style_map[level] 设置 run 级字体和段落行距。

    当 profile 不存在或该 level 无样式数据时不做任何覆写，
    依赖模板 Heading 样式定义的继承。
    """
    if not profile or level not in profile.style_map:
        return
    si = profile.style_map[level]
    if si.font_name or si.font_size_pt or si.east_asia_font:
        _set_run_fonts(
            run,
            latin_font=si.font_name or "Times New Roman",
            east_asia_font=si.east_asia_font,
            size_pt=si.font_size_pt or 12,
            bold=si.bold,
        )
    _set_paragraph_spacing(paragraph, si.line_spacing)


# ── 模板加载与清空 ───────────────────────────────────────

def _enforce_style_outline_levels(doc: Document):
    """确保 Heading 1-6 样式定义中包含显式 w:outlineLvl（WPS 兼容）。"""
    enforced = 0
    for level in range(1, 7):
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = style.element.makeelement(qn("w:pPr"), {})
            style.element.append(pPr)
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = style.element.makeelement(qn("w:outlineLvl"), {})
            pPr.append(outline)
        expected_val = str(level - 1)
        if outline.get(qn("w:val")) != expected_val:
            outline.set(qn("w:val"), expected_val)
            enforced += 1
    if enforced:
        logger.info("已修正 %d 个 Heading 样式的 outlineLvl 定义", enforced)


def load_template(template_path: str) -> Document:
    """加载 Word 模板，清空所有段落内容但保留底层样式定义。"""
    doc = Document(template_path)

    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)

    _enforce_style_outline_levels(doc)
    logger.info("模板加载完成，已清空内容，保留 %d 个样式", len(doc.styles))
    return doc


# ── 核心渲染函数 ─────────────────────────────────────────

def _get_heading_style_name(level: int, profile: HierarchyProfile | None = None) -> str:
    """将 level 映射为 Word Heading 样式名。

    当有 profile 时用其 max_heading_level 做上界；否则 1-6。
    """
    max_lvl = profile.max_heading_level if profile else 6
    clamped = max(1, min(level, max_lvl))
    return f"Heading {clamped}"


def _add_heading_paragraph(
    doc: Document, text: str, level: int,
    profile: HierarchyProfile | None = None,
):
    """添加 title / heading 段落，应用 profile 中学习到的 run 级字体属性。"""
    style_name = _get_heading_style_name(level, profile)
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(_sanitize(text))
    _set_outline_level(para, level)
    _apply_profile_fonts(run, para, level, profile)


def _add_cn_subtitle_paragraph(
    doc: Document, cn_subtitle: str, level: int,
    profile: HierarchyProfile | None = None,
):
    """添加中文小标题段落，应用 profile 中学习到的 run 级字体属性。"""
    style_name = _get_heading_style_name(level, profile)
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(_sanitize(cn_subtitle))
    _set_outline_level(para, level)
    _apply_profile_fonts(run, para, level, profile)


def _add_en_text_paragraph(
    doc: Document, en_text: str,
    profile: HierarchyProfile | None = None,
):
    """添加英文正文段落，字体设置由 profile 驱动。"""
    para = doc.add_paragraph(style="Normal")
    run = para.add_run(_sanitize(en_text))
    _set_outline_level_body(para)

    if profile and profile.en_text_style:
        si = profile.en_text_style
        _set_run_fonts(
            run,
            latin_font=si.font_name or "Times New Roman",
            east_asia_font=si.east_asia_font or "\u5fae\u8f6f\u96c5\u9ed1",
            size_pt=si.font_size_pt or 11,
            bold=si.bold,
        )
    else:
        _set_run_fonts(
            run,
            latin_font="Times New Roman",
            east_asia_font="\u5fae\u8f6f\u96c5\u9ed1",
            size_pt=11,
            bold=False,
        )


def render_blocks_to_docx(
    doc: Document,
    json_blocks: list[dict],
    output_path: str,
    profile: HierarchyProfile | None = None,
):
    """遍历 JSON blocks，按 type/level 动态渲染至 Word 文档并保存。

    参数:
        doc: 已加载并清空的模板 Document 对象
        json_blocks: DeepSeek 返回的结构化 block 列表
        output_path: 输出 .docx 文件路径
        profile: 从模板学习到的层级结构（用于样式验证和约束）
    """
    max_level = profile.max_heading_level if profile else 6
    valid_levels = set(range(1, max_level + 1))
    out_of_range_count = 0

    for i, block in enumerate(json_blocks):
        block_type = block.get("type", "")
        level = block.get("level", 3)

        if level not in valid_levels:
            out_of_range_count += 1
            level = max(1, min(level, max_level))

        if block_type in ("title", "heading"):
            text = block.get("text", "").strip()
            if text:
                _add_heading_paragraph(doc, text, level, profile)

        elif block_type == "annotation":
            cn = block.get("cn_subtitle", "").strip()
            en = block.get("en_text", "").strip()
            if cn:
                _add_cn_subtitle_paragraph(doc, cn, level, profile)
            if en:
                _add_en_text_paragraph(doc, en, profile)

        else:
            text = block.get("text", "") or block.get("en_text", "")
            if text.strip():
                _add_en_text_paragraph(doc, text.strip(), profile)
                logger.warning(
                    "未知 block type '%s'，已作为正文处理 (block #%d)",
                    block_type, i,
                )

    doc.save(output_path)
    logger.info(
        "Word 文档已保存至: %s （共 %d 个 blocks）",
        output_path, len(json_blocks),
    )
    if out_of_range_count:
        logger.warning(
            "%d 个 blocks 的 level 超出模板范围 [1-%d]，已自动修正",
            out_of_range_count, max_level,
        )

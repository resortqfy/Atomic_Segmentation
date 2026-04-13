"""层级分析模块 —— 从示例 docx 文档提取层级结构 Profile，指导下游生成。

工作流程：读取模板文档 → 逐段落提取样式/层级/角色 → 构建 HierarchyProfile。
可作为独立 CLI 运行以调试层级结构。
"""

from __future__ import annotations

import logging
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)

_CN_RE = re.compile(r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef【】]")
_BRACKET_RE = re.compile(r"^【.+?】")


# ── 数据结构 ──────────────────────────────────────────────

@dataclass
class StyleInfo:
    """某一层级的字体/排版信息。"""
    style_name: str
    font_name: str | None = None
    font_size_pt: float | None = None
    bold: bool = False
    east_asia_font: str | None = None
    alignment: str | None = None
    line_spacing: float | None = None
    line_spacing_rule: str | None = None


@dataclass
class HierarchyProfile:
    """从示例文档中学习到的层级结构描述。"""
    max_heading_level: int = 3
    level_roles: dict[int, str] = field(default_factory=dict)
    style_map: dict[int, StyleInfo] = field(default_factory=dict)
    level_sequence_patterns: list[tuple[int, int]] = field(default_factory=list)
    annotation_base_level: int = 3
    cn_subtitle_levels: set[int] = field(default_factory=set)
    en_text_style: StyleInfo = field(
        default_factory=lambda: StyleInfo(
            style_name="Normal",
            font_name="Times New Roman",
            font_size_pt=11,
            bold=False,
            east_asia_font="\u5fae\u8f6f\u96c5\u9ed1",
        )
    )
    level_counts: dict[int, int] = field(default_factory=dict)
    level_samples: dict[int, list[str]] = field(default_factory=dict)

    def describe_for_prompt(self) -> str:
        """生成层级描述文本，可嵌入 DeepSeek system prompt。"""
        lines = ["文档层级结构（从示例文档中学习）："]
        for lvl in sorted(self.level_roles):
            role = self.level_roles[lvl]
            count = self.level_counts.get(lvl, 0)
            lines.append(f"  - level={lvl}: {role}（出现 {count} 次）")
        lines.append(f"  - annotation 基准 level={self.annotation_base_level}")
        if self.cn_subtitle_levels:
            lines.append(
                f"  - 中文小标题使用的 level: {sorted(self.cn_subtitle_levels)}"
            )
        return "\n".join(lines)


# ── 辅助函数 ──────────────────────────────────────────────

def _extract_heading_level(style_name: str) -> int | None:
    """从样式名中提取 Heading 数字，如 'Heading 3' -> 3。"""
    if not style_name:
        return None
    m = re.match(r"[Hh]eading\s*(\d+)", style_name)
    return int(m.group(1)) if m else None


def _is_chinese_text(text: str) -> bool:
    cn_chars = len(_CN_RE.findall(text))
    return cn_chars > len(text) * 0.15 if text else False


def _is_bracket_subtitle(text: str) -> bool:
    return bool(_BRACKET_RE.match(text.strip()))


def _get_run_font_info(paragraph) -> StyleInfo:
    """从段落的第一个 run 提取实际字体信息。"""
    style_name = paragraph.style.name if paragraph.style else "Normal"
    info = StyleInfo(style_name=style_name)

    if not paragraph.runs:
        return _get_style_font_info(paragraph.style, info)

    run = paragraph.runs[0]
    if run.font.name:
        info.font_name = run.font.name
    if run.font.size:
        info.font_size_pt = run.font.size.pt
    info.bold = bool(run.bold)

    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            ea = rFonts.get(qn("w:eastAsia"))
            if ea:
                info.east_asia_font = ea
            if not info.font_name:
                info.font_name = rFonts.get(qn("w:ascii"))

    if not info.font_name or not info.font_size_pt:
        info = _get_style_font_info(paragraph.style, info)

    alignment = paragraph.alignment
    if alignment is not None:
        info.alignment = str(alignment)

    pf = paragraph.paragraph_format
    if pf.line_spacing is not None:
        info.line_spacing = float(pf.line_spacing)
    if pf.line_spacing_rule is not None:
        info.line_spacing_rule = str(pf.line_spacing_rule)

    return info


def _get_style_font_info(style, base: StyleInfo) -> StyleInfo:
    """从样式定义中补充字体信息。"""
    if style is None:
        return base
    try:
        if style.font:
            if not base.font_name and style.font.name:
                base.font_name = style.font.name
            if not base.font_size_pt and style.font.size:
                base.font_size_pt = style.font.size.pt
            if not base.bold:
                base.bold = bool(style.font.bold)
    except Exception:
        pass

    style_elem = style.element if hasattr(style, "element") else None
    if style_elem is not None:
        rPr = style_elem.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                if not base.east_asia_font:
                    base.east_asia_font = rFonts.get(qn("w:eastAsia"))
                if not base.font_name:
                    base.font_name = rFonts.get(qn("w:ascii"))

    if style.base_style:
        base = _get_style_font_info(style.base_style, base)

    return base


def _compute_majority_style(style_list: list[StyleInfo]) -> StyleInfo:
    """从多个 StyleInfo 中按属性取众数，合成一个代表性 StyleInfo。"""
    if not style_list:
        return StyleInfo(style_name="Normal")
    if len(style_list) == 1:
        return style_list[0]

    def _majority(values):
        """取非 None 值的众数；全部为 None 时返回 None。"""
        filtered = [v for v in values if v is not None]
        if not filtered:
            return None
        return Counter(filtered).most_common(1)[0][0]

    def _majority_nullable(values):
        """取包含 None 在内的众数（用于 line_spacing 等可选覆写属性）。"""
        return Counter(values).most_common(1)[0][0]

    def _majority_bool(values):
        return Counter(values).most_common(1)[0][0]

    return StyleInfo(
        style_name=_majority([s.style_name for s in style_list]) or "Normal",
        font_name=_majority([s.font_name for s in style_list]),
        font_size_pt=_majority([s.font_size_pt for s in style_list]),
        bold=_majority_bool([s.bold for s in style_list]),
        east_asia_font=_majority_nullable([s.east_asia_font for s in style_list]),
        alignment=_majority([s.alignment for s in style_list]),
        line_spacing=_majority_nullable([s.line_spacing for s in style_list]),
        line_spacing_rule=_majority_nullable([s.line_spacing_rule for s in style_list]),
    )


# ── 核心分析 ──────────────────────────────────────────────

def analyze_template_hierarchy(docx_path: str) -> HierarchyProfile:
    """读取示例文档，返回 HierarchyProfile。"""
    doc = Document(docx_path)
    profile = HierarchyProfile()

    level_counts: Counter[int] = Counter()
    level_samples: dict[int, list[str]] = defaultdict(list)
    level_styles: dict[int, list[StyleInfo]] = defaultdict(list)
    normal_styles: list[StyleInfo] = []
    transitions: list[tuple[int, int]] = []

    cn_at_level: Counter[int] = Counter()
    heading_at_level: Counter[int] = Counter()

    prev_level: int | None = None
    para_data: list[dict] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        hlevel = _extract_heading_level(style_name)
        font_info = _get_run_font_info(para)

        entry = {
            "text": text,
            "style": style_name,
            "heading_level": hlevel,
            "is_cn": _is_chinese_text(text),
            "is_bracket": _is_bracket_subtitle(text),
            "font_info": font_info,
        }
        para_data.append(entry)

        if hlevel is not None:
            level_counts[hlevel] += 1
            if len(level_samples[hlevel]) < 5:
                level_samples[hlevel].append(text[:80])
            level_styles[hlevel].append(font_info)

            if entry["is_bracket"] or (entry["is_cn"] and hlevel >= 3):
                cn_at_level[hlevel] += 1
            else:
                heading_at_level[hlevel] += 1

            if prev_level is not None:
                transitions.append((prev_level, hlevel))
            prev_level = hlevel
        else:
            if style_name == "Normal":
                normal_styles.append(font_info)
            prev_level = None

    profile.level_counts = dict(level_counts)
    profile.level_samples = dict(level_samples)
    profile.max_heading_level = max(level_counts.keys()) if level_counts else 3

    unique_transitions = list(set(transitions))
    unique_transitions.sort()
    profile.level_sequence_patterns = unique_transitions

    profile.cn_subtitle_levels = {
        lvl for lvl, cnt in cn_at_level.items() if cnt >= 2
    }

    for lvl in sorted(level_counts.keys()):
        if lvl == 1:
            role = "论文主标题 (title)"
        elif lvl == 2:
            role = "一级章节标题 (heading)"
        elif lvl in profile.cn_subtitle_levels:
            cn_ratio = cn_at_level.get(lvl, 0) / max(level_counts[lvl], 1)
            if cn_ratio > 0.5:
                role = "中文小标题 annotation (cn_subtitle)"
            else:
                role = "子章节标题 / 中文小标题 混合 (heading + cn_subtitle)"
        else:
            role = f"{'子' * (lvl - 2)}章节标题 (heading)"
        profile.level_roles[lvl] = role

    annotation_candidates = sorted(profile.cn_subtitle_levels) if profile.cn_subtitle_levels else []
    if annotation_candidates:
        profile.annotation_base_level = annotation_candidates[0]
    else:
        profile.annotation_base_level = min(profile.max_heading_level + 1, 6)

    for lvl, style_list in level_styles.items():
        if style_list:
            profile.style_map[lvl] = _compute_majority_style(style_list)

    if normal_styles:
        profile.en_text_style = _compute_majority_style(normal_styles)

    logger.info(
        "层级分析完成: max_level=%d, annotation_base=%d, cn_levels=%s",
        profile.max_heading_level,
        profile.annotation_base_level,
        sorted(profile.cn_subtitle_levels),
    )
    return profile


# ── CLI 调试入口 ──────────────────────────────────────────

def print_hierarchy_debug(docx_path: str) -> None:
    """打印文档的层级结构调试信息。"""
    print(f"\n{'=' * 70}")
    print(f"  层级分析调试报告: {docx_path}")
    print(f"{'=' * 70}\n")

    doc = Document(docx_path)

    print("── 1. 所有段落样式统计 ──")
    style_counter: Counter[str] = Counter()
    for para in doc.paragraphs:
        if para.text.strip():
            style_counter[para.style.name if para.style else "(None)"] += 1
    for style_name, count in style_counter.most_common():
        print(f"  {style_name:30s}  x{count}")

    print(f"\n── 2. 各 Heading 层级详情 ──")
    profile = analyze_template_hierarchy(docx_path)

    for lvl in sorted(profile.level_counts.keys()):
        count = profile.level_counts[lvl]
        role = profile.level_roles.get(lvl, "未知")
        si = profile.style_map.get(lvl)
        cn_mark = " [CN]" if lvl in profile.cn_subtitle_levels else ""
        print(f"\n  Level {lvl}: {count} 个段落  |  角色: {role}{cn_mark}")
        if si:
            print(f"    字体: {si.font_name or '(继承)'}, "
                  f"大小: {si.font_size_pt or '(继承)'}pt, "
                  f"粗体: {si.bold}, "
                  f"东亚字体: {si.east_asia_font or '(继承)'}")
        samples = profile.level_samples.get(lvl, [])
        for j, s in enumerate(samples[:3]):
            print(f"    样例{j + 1}: {s}")

    print(f"\n── 3. Normal 样式（英文正文）详情 ──")
    si = profile.en_text_style
    print(f"  字体: {si.font_name}, 大小: {si.font_size_pt}pt, "
          f"粗体: {si.bold}, 东亚字体: {si.east_asia_font}")

    print(f"\n── 4. 层级转换模式 ──")
    for parent, child in profile.level_sequence_patterns:
        print(f"  Level {parent} -> Level {child}")

    print(f"\n── 5. Profile 摘要 ──")
    print(f"  最大 Heading 层级: {profile.max_heading_level}")
    print(f"  Annotation 基准层级: {profile.annotation_base_level}")
    print(f"  中文小标题层级: {sorted(profile.cn_subtitle_levels)}")

    print(f"\n── 6. DeepSeek Prompt 描述 ──")
    print(profile.describe_for_prompt())

    print(f"\n{'=' * 70}")
    print("  分析结束")
    print(f"{'=' * 70}\n")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python hierarchy_analyzer.py <docx_path>")
        sys.exit(1)

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    )
    print_hierarchy_debug(sys.argv[1])

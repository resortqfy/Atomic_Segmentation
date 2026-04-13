"""端到端集成测试 —— 验证层级分析、prompt 构建、渲染管线、自适应并发的完整链路。

不依赖外部 API（DeepSeek / Doc2X），通过 mock 数据和本地模板完成测试。

用法:
    python test_integration.py
"""

from __future__ import annotations

import asyncio
import os
import sys
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("test_integration")

TEMPLATE_PATH = "合并_正文标注_20251006_第3届.docx"
OUTPUT_PATH = "test_integration_output.docx"

PASS = 0
FAIL = 0


def check(name: str, condition: bool, detail: str = ""):
    global PASS, FAIL
    if condition:
        PASS += 1
        logger.info("[PASS] %s %s", name, detail)
    else:
        FAIL += 1
        logger.error("[FAIL] %s %s", name, detail)


# ── Test 1: HierarchyProfile 提取 ────────────────────────

def test_hierarchy_analysis():
    logger.info("=" * 60)
    logger.info("Test 1: 层级分析模块")
    logger.info("=" * 60)

    from hierarchy_analyzer import analyze_template_hierarchy, HierarchyProfile

    if not os.path.isfile(TEMPLATE_PATH):
        logger.error("模板文件不存在: %s，跳过此测试", TEMPLATE_PATH)
        return None

    profile = analyze_template_hierarchy(TEMPLATE_PATH)
    check("profile type", isinstance(profile, HierarchyProfile))
    check("max_heading_level >= 1", profile.max_heading_level >= 1,
          f"got {profile.max_heading_level}")
    check("level_roles non-empty", len(profile.level_roles) > 0,
          f"got {len(profile.level_roles)} roles")
    check("level_counts non-empty", len(profile.level_counts) > 0)
    check("annotation_base_level >= 3", profile.annotation_base_level >= 3,
          f"got {profile.annotation_base_level}")
    check("cn_subtitle_levels non-empty", len(profile.cn_subtitle_levels) > 0,
          f"got {sorted(profile.cn_subtitle_levels)}")
    check("en_text_style has font", profile.en_text_style.font_name is not None,
          f"got {profile.en_text_style.font_name}")
    check("describe_for_prompt works", len(profile.describe_for_prompt()) > 50)

    check("level 1 = title", "title" in profile.level_roles.get(1, "").lower())
    check("level 2 = heading", "heading" in profile.level_roles.get(2, "").lower())
    check("level 3 = cn_subtitle", "cn_subtitle" in profile.level_roles.get(3, "").lower())

    return profile


# ── Test 2: System prompt 动态构建 ────────────────────────

def test_system_prompt(profile):
    logger.info("=" * 60)
    logger.info("Test 2: DeepSeek system prompt 构建")
    logger.info("=" * 60)

    from deepseek_client import build_system_prompt

    prompt_with_profile = build_system_prompt(profile)
    prompt_without = build_system_prompt(None)

    check("prompt with profile is str", isinstance(prompt_with_profile, str))
    check("prompt without profile is str", isinstance(prompt_without, str))
    check("profile prompt longer", len(prompt_with_profile) >= len(prompt_without),
          f"{len(prompt_with_profile)} vs {len(prompt_without)}")
    check("profile prompt contains level info",
          "level=3" in prompt_with_profile and "cn_subtitle" in prompt_with_profile)
    check("prompt contains JSON schema", '"type": "annotation"' in prompt_with_profile)


# ── Test 3: Block 校验（profile-aware）────────────────────

def test_validate_blocks(profile):
    logger.info("=" * 60)
    logger.info("Test 3: Block 校验逻辑（含 profile 约束）")
    logger.info("=" * 60)

    from deepseek_client import _validate_blocks

    blocks = [
        {"type": "annotation", "level": 2, "cn_subtitle": "【测试】", "en_text": "test"},
        {"type": "annotation", "level": 3, "cn_subtitle": "【正常】", "en_text": "ok"},
        {"type": "heading", "level": 1, "text": "Title"},
        {"type": "annotation", "level": 99, "cn_subtitle": "【超限】", "en_text": "overflow"},
        {"type": "annotation"},
        {"garbage": True},
    ]

    result = _validate_blocks(blocks, profile)

    check("valid blocks count", len(result) == 5,
          f"expected 5, got {len(result)}")

    anno_levels = [b["level"] for b in result if b["type"] == "annotation"]
    min_anno = profile.annotation_base_level
    max_lvl = profile.max_heading_level
    check("annotation levels >= base",
          all(l >= min_anno for l in anno_levels),
          f"levels={anno_levels}, base={min_anno}")
    check("annotation levels <= max",
          all(l <= max_lvl for l in anno_levels),
          f"levels={anno_levels}, max={max_lvl}")

    result_no_profile = _validate_blocks(
        [{"type": "annotation", "level": 2}], None
    )
    check("no-profile clamp to 3",
          result_no_profile[0]["level"] >= 3)


# ── Test 4: DOCX 渲染（profile-driven）───────────────────

def test_docx_rendering(profile):
    logger.info("=" * 60)
    logger.info("Test 4: DOCX 渲染（含 profile 驱动）")
    logger.info("=" * 60)

    from docx import Document
    from docx_generator import load_template, render_blocks_to_docx

    test_blocks = [
        {"type": "heading", "level": 1, "text": "Test Paper Title"},
        {"type": "heading", "level": 2, "text": "Abstract"},
        {"type": "annotation", "level": 3,
         "cn_subtitle": "【研究背景】", "en_text": "Background information."},
        {"type": "annotation", "level": 4,
         "cn_subtitle": "【步骤一】", "en_text": "Step one description."},
        {"type": "annotation", "level": 5,
         "cn_subtitle": "【深层标注】", "en_text": "Deeper annotation."},
        {"type": "heading", "level": 2, "text": "1. Introduction"},
        {"type": "annotation", "level": 3,
         "cn_subtitle": "【引言概述】", "en_text": "Introduction overview."},
        {"type": "annotation", "level": 7,
         "cn_subtitle": "【超范围level】", "en_text": "Out of range level."},
    ]

    doc = load_template(TEMPLATE_PATH)
    render_blocks_to_docx(doc, test_blocks, OUTPUT_PATH, profile)

    check("output file created", os.path.isfile(OUTPUT_PATH))

    doc_out = Document(OUTPUT_PATH)
    paras = [p for p in doc_out.paragraphs if p.text.strip()]
    check("output has paragraphs", len(paras) > 0, f"got {len(paras)}")

    heading_styles = set()
    for p in paras:
        if p.style and p.style.name.startswith("Heading"):
            heading_styles.add(p.style.name)

    check("output uses Heading 1", "Heading 1" in heading_styles, str(heading_styles))
    check("output uses Heading 2", "Heading 2" in heading_styles, str(heading_styles))
    check("output uses Heading 3", "Heading 3" in heading_styles, str(heading_styles))

    normal_paras = [p for p in paras if p.style and p.style.name == "Normal"]
    check("output has Normal paragraphs", len(normal_paras) > 0,
          f"got {len(normal_paras)}")

    if normal_paras and normal_paras[0].runs:
        run = normal_paras[0].runs[0]
        check("Normal font is TNR", run.font.name == "Times New Roman",
              f"got {run.font.name}")

    for p in paras:
        if p.style and p.style.name.startswith("Heading"):
            from hierarchy_analyzer import _extract_heading_level
            lvl = _extract_heading_level(p.style.name)
            check(f"heading level {lvl} within range",
                  lvl is not None and 1 <= lvl <= profile.max_heading_level,
                  f"style={p.style.name}")

    return True


# ── Test 5: 自适应并发 + 进度追踪 ────────────────────────

def test_adaptive_concurrency():
    logger.info("=" * 60)
    logger.info("Test 5: 自适应并发 + 进度追踪")
    logger.info("=" * 60)

    from main import AdaptiveSemaphore, ProgressTracker
    import config

    async def run_test():
        sem = AdaptiveSemaphore(
            initial=8, min_workers=4, max_workers=16, enabled=True
        )
        tracker = ProgressTracker(total=10)

        for i in range(6):
            await tracker.start_one()
            await sem.acquire()
            await sem.report_latency(2.0, 0)
            sem.release()
            await tracker.complete_one(success=True)

        check("concurrency increased (fast)", sem.current_limit > 8,
              f"got {sem.current_limit}")

        for i in range(6):
            await tracker.start_one()
            await sem.acquire()
            await sem.report_latency(40.0, 5)
            sem.release()
            await tracker.complete_one(success=True)

        check("concurrency decreased (slow)", sem.current_limit < 16,
              f"got {sem.current_limit}")

        check("tracker total correct", tracker.total == 10)

    asyncio.run(run_test())

    sem_disabled = asyncio.run(_test_disabled())
    check("disabled semaphore stays at initial", sem_disabled == 8,
          f"got {sem_disabled}")


async def _test_disabled():
    from main import AdaptiveSemaphore
    sem = AdaptiveSemaphore(initial=8, min_workers=4, max_workers=16, enabled=False)
    for _ in range(10):
        await sem.report_latency(1.0, 0)
    return sem.current_limit


# ── Test 6: Markdown 分块 + backmatter 过滤 ──────────────

def test_markdown_splitting():
    logger.info("=" * 60)
    logger.info("Test 6: Markdown 分块 + backmatter 过滤")
    logger.info("=" * 60)

    from main import _split_markdown_into_sections, _is_backmatter

    md = """# Paper Title

Some preamble text.

## Abstract

This is the abstract of the paper. It contains important information.

## 1. Introduction

This section introduces the topic.

### 1.1 Background

Background details here.

## References

[1] Some reference.
"""

    sections = _split_markdown_into_sections(md)
    check("sections found", len(sections) > 3, f"got {len(sections)}")

    titles = [s["title"] for s in sections]
    check("Paper Title in sections", "Paper Title" in titles)
    check("Abstract in sections", "Abstract" in titles)
    check("heading_level correct for ##", 
          any(s["heading_level"] == 2 for s in sections if s["title"] == "Abstract"))
    check("heading_level correct for ###",
          any(s["heading_level"] == 3 for s in sections if "Background" in s["title"]))

    check("References is backmatter", _is_backmatter("References"))
    check("5. References is backmatter", _is_backmatter("5. References"))
    check("Introduction is NOT backmatter", not _is_backmatter("1. Introduction"))
    check("Acknowledgements is backmatter", _is_backmatter("Acknowledgements"))


# ── Test 7: 渲染输出层级与模板对比验证 ────────────────────

def test_output_hierarchy_matches_template(profile):
    logger.info("=" * 60)
    logger.info("Test 7: 输出层级与模板一致性验证")
    logger.info("=" * 60)

    if not os.path.isfile(OUTPUT_PATH):
        logger.warning("输出文件不存在，跳过此测试")
        return

    from docx import Document
    from hierarchy_analyzer import _extract_heading_level

    doc = Document(OUTPUT_PATH)

    output_levels = set()
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p.style and p.style.name.startswith("Heading"):
            lvl = _extract_heading_level(p.style.name)
            if lvl is not None:
                output_levels.add(lvl)

    template_levels = set(profile.level_counts.keys())

    check("output levels subset of template levels",
          output_levels.issubset(template_levels),
          f"output={sorted(output_levels)}, template={sorted(template_levels)}")

    check("no level > max_heading_level",
          all(l <= profile.max_heading_level for l in output_levels),
          f"output levels={sorted(output_levels)}, max={profile.max_heading_level}")

    check("output uses annotation-range levels",
          any(l >= profile.annotation_base_level for l in output_levels),
          f"output levels={sorted(output_levels)}, base={profile.annotation_base_level}")


# ── Test 8: outlineLvl 段落级验证（WPS 兼容）─────────────

def test_outline_levels(profile):
    logger.info("=" * 60)
    logger.info("Test 8: outlineLvl 段落级写入验证（WPS 兼容）")
    logger.info("=" * 60)

    from docx import Document
    from docx.oxml.ns import qn
    from docx_generator import load_template, render_blocks_to_docx
    from hierarchy_analyzer import _extract_heading_level

    test_blocks = [
        {"type": "heading", "level": 1, "text": "OutlineLvl Test Title"},
        {"type": "heading", "level": 2, "text": "Abstract"},
        {"type": "annotation", "level": 3,
         "cn_subtitle": "【测试标注】", "en_text": "Test English text."},
        {"type": "annotation", "level": 4,
         "cn_subtitle": "【子级标注】", "en_text": "Sub annotation."},
        {"type": "heading", "level": 2, "text": "Introduction"},
        {"type": "annotation", "level": 5,
         "cn_subtitle": "【深层标注】", "en_text": "Deep level."},
    ]

    outline_test_path = "test_outline_output.docx"
    doc = load_template(TEMPLATE_PATH)
    render_blocks_to_docx(doc, test_blocks, outline_test_path, profile)

    doc_out = Document(outline_test_path)

    heading_paras_checked = 0
    heading_paras_ok = 0
    normal_paras_checked = 0
    normal_paras_ok = 0

    for p in doc_out.paragraphs:
        if not p.text.strip():
            continue

        pPr = p._element.find(qn("w:pPr"))
        style_name = p.style.name if p.style else "Normal"
        hlevel = _extract_heading_level(style_name)

        if hlevel is not None:
            heading_paras_checked += 1
            if pPr is not None:
                outline_elem = pPr.find(qn("w:outlineLvl"))
                if outline_elem is not None:
                    val = outline_elem.get(qn("w:val"))
                    expected = str(hlevel - 1)
                    if val == expected:
                        heading_paras_ok += 1
                    else:
                        logger.error(
                            "  outlineLvl 值错误: style=%s, expected=%s, got=%s, text=%s",
                            style_name, expected, val, p.text[:40],
                        )
                else:
                    logger.error(
                        "  缺少 outlineLvl: style=%s, text=%s",
                        style_name, p.text[:40],
                    )
        elif style_name == "Normal":
            normal_paras_checked += 1
            if pPr is not None:
                outline_elem = pPr.find(qn("w:outlineLvl"))
                if outline_elem is not None and outline_elem.get(qn("w:val")) == "9":
                    normal_paras_ok += 1

    check("all heading paras have correct outlineLvl",
          heading_paras_checked > 0 and heading_paras_ok == heading_paras_checked,
          f"{heading_paras_ok}/{heading_paras_checked}")
    check("all Normal paras have outlineLvl=9",
          normal_paras_checked > 0 and normal_paras_ok == normal_paras_checked,
          f"{normal_paras_ok}/{normal_paras_checked}")

    style_h1 = doc_out.styles["Heading 1"]
    pPr_h1 = style_h1.element.find(qn("w:pPr"))
    outline_h1 = pPr_h1.find(qn("w:outlineLvl")) if pPr_h1 is not None else None
    check("Heading 1 style has outlineLvl=0",
          outline_h1 is not None and outline_h1.get(qn("w:val")) == "0")

    style_h3 = doc_out.styles["Heading 3"]
    pPr_h3 = style_h3.element.find(qn("w:pPr"))
    outline_h3 = pPr_h3.find(qn("w:outlineLvl")) if pPr_h3 is not None else None
    check("Heading 3 style has outlineLvl=2",
          outline_h3 is not None and outline_h3.get(qn("w:val")) == "2")

    if os.path.isfile(outline_test_path):
        os.remove(outline_test_path)

    return True


# ── 主入口 ───────────────────────────────────────────────

def main():
    logger.info("\n" + "=" * 70)
    logger.info("  集成测试开始")
    logger.info("=" * 70 + "\n")

    profile = test_hierarchy_analysis()
    if profile is None:
        logger.error("层级分析失败，无法继续后续测试")
        sys.exit(1)

    test_system_prompt(profile)
    test_validate_blocks(profile)
    test_docx_rendering(profile)
    test_adaptive_concurrency()
    test_markdown_splitting()
    test_output_hierarchy_matches_template(profile)
    test_outline_levels(profile)

    logger.info("\n" + "=" * 70)
    logger.info("  测试结果: %d PASS, %d FAIL", PASS, FAIL)
    logger.info("=" * 70 + "\n")

    if os.path.isfile(OUTPUT_PATH):
        os.remove(OUTPUT_PATH)
        logger.info("已清理临时文件: %s", OUTPUT_PATH)

    sys.exit(0 if FAIL == 0 else 1)


if __name__ == "__main__":
    main()
